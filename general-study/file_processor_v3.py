#####################################################################################################
#################### This file processor is design to work with custom models #######################
#####################################################################################################

from pathlib import Path
from docx import Document
from dataclasses import dataclass

from ollama_client import LLMClient
from txt_logger import setup_txt_logger

from pathlib import Path

# Setup json logger
log_file_path = Path(__file__).parent / "file_processor_v2_log.txt"
txt_logger = setup_txt_logger(__name__, log_file_path)


@dataclass(slots=True)
class ProcessFilesConfig:
    files_text: list[str]
    model_name: str
    client: LLMClient
    text_topic: str | None
    output_file_name: str | None


def process_files(config: ProcessFilesConfig) -> None:
    """
    Process files using a LLM.

    Args:
        A ProcessFilesConfig objet with all configurations.
    """
    
    for file_path, text in config.files_text:
        if file_path == "TEMA.docx":
            continue  
        messages: list[dict[str, str]] = []
        messages.append({"role": "system", "content": config.text_topic or ""})
        messages.append({"role": "user", "content": text})

        txt_logger.info({"variable": "model name", "value": config.model_name})
        txt_logger.info({"variable": "config.system_role", "value": config.text_topic})
        txt_logger.info({"variable": "text", "value": text})

        response: str = config.client.chat(
            model_name=config.model_name,
            messages=messages,
            options=config.client.options,
        )

        txt_logger.info({"variable": "response", "value": response})

        if config.output_file_name:
            save_response_to_file(response, file_path, config.output_file_name)
        else:
            raise ValueError("Output file name is not valid.")

##### Helpers functions #####
def load_file_texts_from_folder(folder_path: Path) ->  list[tuple[str, str]]:
    results: list[tuple[str, str]] = []

    for path in folder_path.iterdir():
        if path.is_file() and path.suffix.lower() == ".docx":
            document: Document = Document(path)

            paragraphs_text: list[str] = []
            consecutive_empty_count = 0

            for paragraph in document.paragraphs:
                is_empty = (paragraph.text.strip() == "")

                if is_empty:
                    consecutive_empty_count += 1
                else:
                    consecutive_empty_count = 0

                if consecutive_empty_count == 2:
                    # Stop reading further, excluding these empty paragraphs
                    break

                paragraphs_text.append(paragraph.text)

            content = "\n".join(paragraphs_text)
            results.append((path.name, content))

    return results
    

def load_topic_from_file(file_path: Path) -> str:
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    document = Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    topic = "\n".join(paragraphs)

    return topic


def load_system_role_from_file(file_path: Path) -> str:
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    with file_path.open("r", encoding="utf-8") as f:
        system_role: str = f.read()
        system_role += "\n"

    return system_role

def save_response_to_file(
    response: str,
    source_file_name: str,
    output_file: Path = Path("classification_results.txt")
) -> None:
    output_file = Path.cwd() / output_file
    with output_file.open("a", encoding="utf-8") as f:
        f.write(f"{source_file_name} | {response}")
##### End of helpers functions #####


def main() -> None:
    pass


if __name__ == "__main__":
    main()
