import os
from pathlib import Path
from docx import Document


##### Helpers functions #####

def load_text_from_docx_file(file_path: Path) -> str:
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_path.is_file() and file_path.suffix.lower() == ".docx":
        document: Document = Document(file_path)

        paragraphs_text: list[str] = []
        consecutive_empty_count = 0

        for paragraph in document.paragraphs:
            is_empty = paragraph.text.strip() == ""

            if is_empty:
                consecutive_empty_count += 1
            else:
                consecutive_empty_count = 0

            if consecutive_empty_count == 2:
                # Stop reading further, excluding these empty paragraphs
                break

            paragraphs_text.append(paragraph.text)

        content = "\n".join(paragraphs_text)

    return content

def load_text_from_txt_file(file_path: Path) -> str:
    if not file_path.exists() or not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if file_path.suffix.lower() == ".txt":
        with file_path.open("r", encoding="utf-8") as f:
            content: str = f.read()
            content += "\n"
    
    return content


def save_response_to_file(
    response: str,
    source_file_name: str,
    output_file_path: Path,
) -> None:    
    with output_file_path.open("a", encoding="utf-8") as f:
        f.write(f"{source_file_name},{response}\n")


def remove_blank_lines_from_txt_file(file_path: str) -> None:
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    non_blank_lines: list[str] = [line for line in lines if line.strip()]

    with open(file_path, "w", encoding="utf-8") as f:
        f.writelines(non_blank_lines)

def get_ignored_file_text(file_path_list: list[Path]) -> str:
    ignored_filename = os.getenv("IGNORED_FILE_NAME")
    if not ignored_filename:
        raise ValueError("IGNORED_FILE_NAME não configurado no .env")

    for file_path in file_path_list:
        if file_path.name == ignored_filename:
            if file_path.suffix.lower() == ".docx":
                document = Document(file_path)
                paragraphs = [p.text for p in document.paragraphs]
                return "\n".join(paragraphs)

    raise FileNotFoundError(f"Arquivo ignorado '{ignored_filename}' não encontrado na lista")

def remove_ignored_file_from_path_list(
    file_path_list: list[Path], ignored_filename: str = os.getenv("IGNORED_FILE_NAME")
) -> list[Path]:
    if not ignored_filename:
        raise ValueError("IGNORED_FILE_NAME não configurado no .env")

    return [f for f in file_path_list if f.name != ignored_filename]

# this one can't be replaced since it loads integer labels specifically from docx files
def load_label_from_docx_file(file_path: Path) -> list[int]:
    doc: Document = Document(file_path)
    label_values: list[int] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # check if the text is a digit
        if text.isdigit():
            label_values.append(int(text))

    return label_values


def clear_txt_file(file_path: Path) -> None:
    if not file_path.exists() and not file_path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")
    file_path.open("w", encoding="utf-8").close()
    
# this function needs to be rethought since the new methodology development 
def calculate_averages_per_index(values: list[list[int]]) -> list[float]:
    if not values:
        return []

    list_length: int = len(values)
    index_amount: int = len(values[0])

    sum_per_index: list[int] = [0] * index_amount

    for sublist in values:
        for index, value in enumerate(sublist):
            sum_per_index[index] += value

    averages: list[float] = [
        total / list_length for total in sum_per_index
    ]

    return averages
    

##### End of helpers functions #####


def main() -> None:
    pass


if __name__ == "__main__":
    main()
